#!/usr/bin/env python3
# Requirements: pip install langchain langchain-community langchain-huggingface sentence-transformers faiss-cpu pypdf transformers Pillow scikit-learn nltk tqdm rouge-score bert-score
from anthropic import Anthropic
from dotenv import load_dotenv
import os
import argparse
import base64
import time
import re
import json
import logging
from typing import List, Dict, Any, Tuple, Optional, Union, Callable
import nltk
# Removed: import spacy
from langchain_community.document_loaders import PyPDFLoader, UnstructuredFileLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import UnstructuredPDFLoader


from langchain_community.vectorstores import FAISS
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.llms import HuggingFaceHub
from langchain.chains import LLMChain
from langchain.prompts import PromptTemplate
from langchain_community.cache import InMemoryCache
import langchain
from sentence_transformers import CrossEncoder
from sklearn.feature_extraction.text import TfidfVectorizer
from transformers import pipeline, AutoProcessor, AutoModelForCausalLM, AutoTokenizer
from PIL import Image
import torch
import tqdm


# Configure logging
logging.basicConfig(
   level=logging.INFO,
   format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
   handlers=[
       logging.FileHandler("rag_system.log"),
       logging.StreamHandler()
   ]
)
logger = logging.getLogger("multimodal_rag")


# Set environment variable to avoid parallelism warnings with tokenizers
os.environ["TOKENIZERS_PARALLELISM"] = "false"


# Import your existing prompts
# Ensure 'prompts.py' exists with SYSTEM_MESSAGE and USER_PROMPT defined
try:
   from prompts import SYSTEM_MESSAGE, USER_PROMPT
except ImportError:
   logger.warning("Could not import prompts.py. Using default prompt templates.")


# Initialize cache for embeddings
#langchain.llm_cache = InMemoryCache()


# Download necessary NLTK resources
try:
   nltk.download('punkt', quiet=True)
   nltk.download('stopwords', quiet=True)
   # Removed: spacy.load('en_core_web_sm')
except Exception as e:
   logger.warning(f"Could not download NLTK resources: {e}")




class DocumentProcessor:
   """Handles document loading and preprocessing for different file types."""


   def __init__(self):
       self.supported_extensions = {
           '.pdf': self._load_pdf,
           '.txt': self._load_text,
           '.docx': self._load_docx,
           '.md': self._load_text,
           '.html': self._load_text
       }


   def load_document(self, file_path: str) -> List[Any]:
       """Load a document based on its file extension."""
       _, ext = os.path.splitext(file_path.lower())
       if ext not in self.supported_extensions:
           raise ValueError(f"Unsupported file type: {ext}")


       return self.supported_extensions[ext](file_path)


   def _load_pdf(self, file_path: str) -> List[Any]:
       """Load a PDF file with enhanced metadata."""
       logger.info(f"Loading PDF: {file_path}")
       loader = UnstructuredPDFLoader(file_path)
       documents = loader.load()


       # Enhance metadata with more detailed information
       for i, doc in enumerate(documents):
           # Basic metadata
           doc.metadata["page_number"] = i + 1
           doc.metadata["source"] = file_path
           doc.metadata["total_pages"] = len(documents)


           # Extract more semantic metadata
           # Attempt to identify section titles/headers from content
           header_match = re.search(r'^(.*?)(?:\n|$)', doc.page_content.strip())
           if header_match:
               potential_header = header_match.group(1).strip()
               # If it's less than 10 words, likely a header
               if len(potential_header.split()) < 10 and len(potential_header) < 100:
                   doc.metadata["page_header"] = potential_header


           # Try to identify if page contains figures/tables
           has_figure = bool(re.search(r'fig(ure)?\.?\s*\d+|figure|chart|graph|diagram',
                                       doc.page_content, re.IGNORECASE))
           has_table = bool(re.search(r'table\.?\s*\d+|tabular data',
                                      doc.page_content, re.IGNORECASE))


           doc.metadata["has_figure"] = has_figure
           doc.metadata["has_table"] = has_table


           # Identify page type (TOC, index, content)
           if i < 3:  # First few pages
               if re.search(r'contents|table of contents|toc', doc.page_content, re.IGNORECASE):
                   doc.metadata["page_type"] = "table_of_contents"


           # Check if it's an index page
           if re.search(r'index|glossary', doc.page_content, re.IGNORECASE):
               doc.metadata["page_type"] = "index"


       return documents


   def _load_text(self, file_path: str) -> List[Any]:
       """Load a text file (txt, md, html) with enhanced section detection."""
       logger.info(f"Loading text file: {file_path}")
       loader = UnstructuredFileLoader(file_path)
       documents = loader.load()


       # Process text files to identify sections and enhance metadata
       for doc in documents:
           doc.metadata["source"] = file_path


           # Split content into lines to identify structure
           lines = doc.page_content.split('\n')


           # Identify headers/sections by looking for patterns
           headers = []
           for i, line in enumerate(lines):
               # Check for markdown headers
               if re.match(r'^#{1,6}\s+', line):
                   headers.append((i, line.strip()))
               # Check for underlined headers
               elif i > 0 and re.match(r'^[=\-]+$', line) and lines[i - 1].strip():
                   headers.append((i - 1, lines[i - 1].strip()))


           # Add headers to metadata
           if headers:
               doc.metadata["sections"] = [h[1] for h in headers]


           # Try to identify document type based on content
           if re.search(r'README', file_path, re.IGNORECASE) or \
                   (lines and re.match(r'^#{1,2}\s+', lines[0])):
               doc.metadata["doc_type"] = "readme"
           elif re.search(r'license|copyright|mit|apache|gnu',
                          doc.page_content, re.IGNORECASE):
               doc.metadata["doc_type"] = "license"
           elif re.search(r'changelog|release notes|version history',
                          doc.page_content, re.IGNORECASE):
               doc.metadata["doc_type"] = "changelog"


       return documents


   def _load_docx(self, file_path: str) -> List[Any]:
       """Load a DOCX file with style and formatting metadata."""
       logger.info(f"Loading DOCX: {file_path}")
       loader = UnstructuredFileLoader(file_path)
       documents = loader.load()


       for doc in documents:
           doc.metadata["source"] = file_path


           # Additional processing for DOCX if using python-docx for more metadata
           try:
               import docx
               doc_obj = docx.Document(file_path)


               # Extract styles and formatting information
               styles = set()
               for paragraph in doc_obj.paragraphs:
                   if paragraph.style.name:
                       styles.add(paragraph.style.name)


               # Identify headers based on styles
               headers = []
               for para in doc_obj.paragraphs:
                   if para.style.name and "Heading" in para.style.name:
                       headers.append(para.text)


               if headers:
                   doc.metadata["headers"] = headers
               if styles:
                   doc.metadata["styles"] = list(styles)


           except ImportError:
               logger.warning("python-docx not installed. Enhanced DOCX metadata extraction unavailable.")
           except Exception as e:
               logger.error(f"Error extracting DOCX metadata: {e}")


       return documents




class AdaptiveTextSplitter:
   """Text splitter that adapts to document structure and type with improved chunking."""


   def __init__(self):
       # Different splitting strategies
       # Removed 'spacy' splitter
       self.splitters = {
           "recursive": RecursiveCharacterTextSplitter
           #"layout": LayoutTextSplitter,
       }


       # Configurations for different document types
       self.document_configs = {
           "pdf": {
               "splitter": "recursive", # Use layout for PDFs
               "chunk_size": 1000,
               "chunk_overlap": 200
           },
           "text": {
               "splitter": "recursive", # Use recursive for text
               "chunk_size": 800,
               "chunk_overlap": 150
           },
            "docx": { # Added docx config
               "splitter": "recursive",
               "chunk_size": 900,
               "chunk_overlap": 180
           },
           "default": {
               "splitter": "recursive", # Default to recursive
               "chunk_size": 800,
               "chunk_overlap": 200
           }
       }


   def analyze_document(self, document: List[Any]) -> Dict[str, Any]:
       """Analyze document to determine optimal chunking parameters."""
       # Count sections using regex patterns for headers
       section_pattern = r'^\s*#{1,6}\s+.+|^[A-Z][^.!?]*[.!?]$'
       section_matches = [
           re.findall(section_pattern, doc.page_content, re.MULTILINE)
           for doc in document
       ]
       total_sections = sum(len(matches) for matches in section_matches)


       # Analyze average paragraph length
       paragraph_splits = [
           re.split(r'\n\s*\n', doc.page_content)
           for doc in document
       ]
       paragraphs = [p for splits in paragraph_splits for p in splits if p.strip()]


       total_paragraphs = len(paragraphs)
       # Handle division by zero if no paragraphs are found
       avg_paragraph_length = sum(len(p.split()) for p in paragraphs) / max(1, total_paragraphs)


       # Analyze the presence of structured elements
       has_lists = any(
           bool(re.search(r'^\s*[-*•]\s+', doc.page_content, re.MULTILINE))
           for doc in document
       )


       has_code_blocks = any(
           bool(re.search(r'```|^\s{4}', doc.page_content, re.MULTILINE))
           for doc in document
       )


       # Analyze semantic coherence - using NLTK sentence tokenizer
       sentences = []
       try:
           for doc in document:
               sentences.extend(nltk.sent_tokenize(doc.page_content))
       except LookupError:
           logger.warning("NLTK 'punkt' tokenizer not found. Sentence analysis skipped.")




       # Determine optimal parameters based on document analysis
       config = self.document_configs["default"].copy()


       # Logic for determining optimal chunk size and overlap
       if total_sections > 0:
           # Document with clear section structure
           section_density = total_sections / len(document) if len(document) > 0 else 0


           if section_density > 0.5:  # Many sections per page
               # Use smaller chunks to preserve section boundaries
               config["chunk_size"] = 600
               config["chunk_overlap"] = 150
           else:  # Fewer sections
               # Use larger chunks to keep sections together
               config["chunk_size"] = 1200
               config["chunk_overlap"] = 300


       elif avg_paragraph_length > 100:  # Very long paragraphs
           # Break up long paragraphs with smaller chunks and more overlap
           config["chunk_size"] = 700
           config["chunk_overlap"] = 250


       elif avg_paragraph_length < 30 and total_paragraphs > 20:  # Many short paragraphs
           # Use larger chunks to group related short paragraphs
           config["chunk_size"] = 1500
           config["chunk_overlap"] = 400


       # Special handling for certain content types
       if has_code_blocks:
           # Preserve code blocks by using larger chunks
           config["chunk_size"] = max(config["chunk_size"], 1500)
           config["chunk_overlap"] = max(config["chunk_overlap"], 300)


       if has_lists:
           # Ensure lists aren't broken up
           config["chunk_overlap"] = max(config["chunk_overlap"], 300)


       logger.info(f"Document analysis results: sections={total_sections}, avg_para_len={avg_paragraph_length:.1f}, "
                   f"has_lists={has_lists}, has_code={has_code_blocks}")
       logger.info(f"Selected chunk_size={config['chunk_size']}, chunk_overlap={config['chunk_overlap']}")


       return config


   def get_splitter_for_document(self, document: List[Any], doc_type: str = "default") -> Any:
       """Get the appropriate text splitter for the document."""
       # Get base configuration
       config = self.document_configs.get(doc_type, self.document_configs["default"])


       # Adapt config based on document analysis
       analysis_config = self.analyze_document(document)
       config.update(analysis_config)


       # Create the splitter
       splitter_type = config.get("splitter", "recursive") # Default to recursive if not specified
       splitter_class = self.splitters.get(splitter_type, RecursiveCharacterTextSplitter) # Fallback to Recursive


       # RecursiveCharacterTextSplitter or fallback case
       return splitter_class(
           chunk_size=config["chunk_size"],
           chunk_overlap=config["chunk_overlap"],
           separators=["\n\n\n", "\n\n", "\n", ". ", " ", ""] # Common separators
       )




   def split_documents(self, documents: List[Any], doc_type: str = "default") -> List[Any]:
       """Split documents using the appropriate strategy with enhanced metadata."""
       start_time = time.time()


       splitter = self.get_splitter_for_document(documents, doc_type)
       chunks = splitter.split_documents(documents)


       # Process chunks to improve quality and add metadata
       enhanced_chunks = []


       # Group chunks that likely belong together
       i = 0
       while i < len(chunks):
           current_chunk = chunks[i]


           # Initialize metadata
           current_chunk.metadata["chunk_id"] = i
           current_chunk.metadata["total_chunks"] = len(chunks)


           # Extract potential headers - simplified pattern
           header_match = re.search(r'^\s*(#{1,6}\s+.+)',
                                    current_chunk.page_content.strip()[:200], re.MULTILINE)
           if header_match:
               current_chunk.metadata["header"] = header_match.group(1).strip()


           #else:
                # Try to capture section titles if using LayoutTextSplitter (often included at chunk start)
               #if isinstance(splitter, LayoutTextSplitter):
                    #layout_header_match = re.search(r'^\s*([A-Za-z0-9\s]{10,100})\s*\n', current_chunk.page_content.strip())
                    #if layout_header_match and len(layout_header_match.group(1).split()) < 15:
                         #current_chunk.metadata["potential_section_title"] = layout_header_match.group(1).strip()




           # Skip empty or nearly empty chunks
           if len(current_chunk.page_content.strip()) < 50:
               i += 1
               continue


           # Look for incomplete sentences at the end of a chunk using NLTK
           try:
               sentences = nltk.sent_tokenize(current_chunk.page_content.strip())
               if sentences:
                   last_sentence = sentences[-1]
                   # Check if the last sentence looks incomplete
                   if not re.search(r'[.!?]\s*$', last_sentence) and i < len(chunks) - 1:
                        # Check if we should merge with next chunk
                       next_chunk = chunks[i + 1]


                       # Calculate the total length if merged
                       merged_length = len(current_chunk.page_content) + len(next_chunk.page_content)


                       # If merging wouldn't create too large a chunk, combine them
                       # Allow slightly larger chunks than the target size
                       if merged_length <= splitter.chunk_size * 1.3:
                           logger.debug(f"Merging chunk {i} and {i+1} due to incomplete sentence.")
                           current_chunk.page_content += "\n" + next_chunk.page_content


                           # Merge metadata from both chunks (prioritize current chunk's metadata)
                           for k, v in next_chunk.metadata.items():
                               if k not in current_chunk.metadata:
                                   current_chunk.metadata[k] = v
                               else:
                                   # Handle potential list metadata conflicts by combining
                                   if isinstance(current_chunk.metadata[k], list) and isinstance(v, list):
                                        current_chunk.metadata[k] = list(set(current_chunk.metadata[k] + v))
                                   # Simple overwrite for others - could be improved
                                   # pass


                           # Skip the next chunk since we merged it
                           i += 2
                           enhanced_chunks.append(current_chunk)
                           continue # Continue the while loop with the new 'i'
           except LookupError:
               logger.warning("NLTK 'punkt' tokenizer not available for sentence-based merging.")
           except Exception as e:
               logger.error(f"Error during sentence analysis and merging: {e}")




           # Add semantic labels based on content
           content = current_chunk.page_content.lower()


           if re.search(r'(table|figure|chart|graph)\s+\d+', content):
               current_chunk.metadata["contains_visual_element"] = True


           if re.search(r'(definition|defined as|refers to)', content):
               current_chunk.metadata["contains_definition"] = True


           if re.search(r'(in conclusion|to summarize|in summary)', content):
               current_chunk.metadata["contains_conclusion"] = True


           if re.search(r'(step|procedure|process|how to)', content):
               current_chunk.metadata["contains_procedure"] = True


           enhanced_chunks.append(current_chunk)
           i += 1 # Move to the next chunk if no merge happened


       end_time = time.time()
       logger.info(
           f"Split {len(documents)} documents into {len(enhanced_chunks)} enhanced chunks "
           f"in {end_time - start_time:.2f} seconds")


       return enhanced_chunks




class ImageProcessor:
   """Process images for captioning and analysis with enhanced features."""


   def __init__(self):
       self.device = "cuda" if torch.cuda.is_available() else "cpu"
       logger.info(f"Image processor using device: {self.device}")


       # Initialize base image captioning model
       try:
           self.captioner = pipeline("image-to-text", model="Salesforce/blip-image-captioning-large",
                                     device=self.device)
           logger.info("Initialized enhanced BLIP image captioning model")
       except Exception as e:
           logger.error(f"Failed to load image captioning model: {e}")
           # Fallback to a smaller model if the large one fails
           try:
               self.captioner = pipeline("image-to-text", model="Salesforce/blip-image-captioning-base",
                                         device=self.device)
               logger.info("Initialized fallback BLIP base image captioning model")
           except Exception as e2:
               logger.error(f"Failed to load fallback captioning model: {e2}")
               self.captioner = None


       # Initialize enhanced object detection for better image analysis
       try:
           self.object_detector = pipeline("object-detection", model="facebook/detr-resnet-50",
                                           device=self.device)
           logger.info("Initialized DETR object detection model")
       except Exception as e:
           logger.warning(f"Failed to load object detection model: {e}")
           self.object_detector = None


       # Initialize scene classification model
       try:
           self.scene_classifier = pipeline("image-classification", model="microsoft/resnet-50",
                                            top_k=5, device=self.device)
           logger.info("Initialized scene classification model")
       except Exception as e:
           logger.warning(f"Failed to load scene classification model: {e}")
           self.scene_classifier = None


       # Initialize OCR capability for text in images
       try:
           import pytesseract
           self.has_ocr = True
           logger.info("OCR capability initialized")
       except ImportError:
           logger.warning("pytesseract not installed, OCR unavailable")
           self.has_ocr = False


   def generate_caption(self, image_path: str) -> str:
       """Generate a detailed caption for the image."""
       if not self.captioner:
           return "An image (no captioning model available)"


       try:
           image = Image.open(image_path)


           # Generate multiple captions with different parameters for diversity
           result1 = self.captioner(image, max_new_tokens=50)
           # Try with different parameters for a second opinion
           result2 = self.captioner(image, max_new_tokens=30, num_beams=5)


           caption1 = result1[0]["generated_text"]
           caption2 = result2[0]["generated_text"] if result2 else ""


           # Combine or select the better caption
           if caption2 and len(caption2) > len(caption1) * 0.7:  # If second caption is substantial
               # Combine unique information from both
               combined = caption1
               # Add unique information from caption2 if not already in caption1
               unique_info = [phrase for phrase in re.split(r'[,.]', caption2)
                              if phrase.strip() and phrase.strip().lower() not in caption1.lower()]


               if unique_info:
                   combined += ". " + " ".join(unique_info)


               logger.info(f"Generated enhanced caption: {combined}")
               return combined
           else:
               logger.info(f"Generated caption: {caption1}")
               return caption1


       except Exception as e:
           logger.error(f"Error generating caption: {e}")
           return "An image (failed to generate caption)"


   def extract_text_from_image(self, image_path: str) -> str:
       """Extract text from image using OCR."""
       if not self.has_ocr:
           return ""


       try:
           import pytesseract
           image = Image.open(image_path)


           # Convert to RGB if needed (for PNG with transparency)
           if image.mode == 'RGBA':
               image = image.convert('RGB')


           # Process image to enhance text recognition
           # For better OCR, we could add preprocessing here


           # Extract text
           text = pytesseract.image_to_string(image)


           if text.strip():
               logger.info(f"Extracted {len(text.split())} words from image via OCR")
               return text.strip()
           return ""
       except Exception as e:
           logger.error(f"Error in OCR: {e}")
           return ""


   def detect_objects(self, image_path: str) -> List[Dict[str, Any]]:
       """Detect objects in the image."""
       if not self.object_detector:
           return []


       try:
           image = Image.open(image_path)
           results = self.object_detector(image)


           # Filter results with reasonable confidence
           filtered_results = [r for r in results if r['score'] > 0.3]


           logger.info(f"Detected {len(filtered_results)} objects in image")
           return filtered_results
       except Exception as e:
           logger.error(f"Error detecting objects: {e}")
           return []


   def classify_scene(self, image_path: str) -> List[Dict[str, float]]:
       """Classify the scene in the image."""
       if not self.scene_classifier:
           return []


       try:
           image = Image.open(image_path)
           results = self.scene_classifier(image)


           logger.info(f"Classified image scene: {results[0]['label']} ({results[0]['score']:.2f})")
           return results
       except Exception as e:
           logger.error(f"Error classifying scene: {e}")
           return []


   def analyze_color_palette(self, image_path: str, num_colors: int = 5) -> List[str]:
       """Extract dominant colors from the image."""
       try:
           from sklearn.cluster import KMeans
           import numpy as np


           image = Image.open(image_path)
           # Resize for faster processing
           image = image.resize((150, 150))
           image_array = np.array(image)


           # Reshape for clustering
           pixels = image_array.reshape(-1, 3)


           # Cluster to find dominant colors
           kmeans = KMeans(n_clusters=num_colors)
           kmeans.fit(pixels)


           # Get the colors
           colors = kmeans.cluster_centers_.astype(int)


           # Convert to hex format
           hex_colors = ['#%02x%02x%02x' % (r, g, b) for r, g, b in colors]


           logger.info(f"Extracted color palette with {len(hex_colors)} colors")
           return hex_colors
       except Exception as e:
           logger.error(f"Error analyzing color palette: {e}")
           return []


   def analyze_image(self, image_path: str) -> Dict[str, Any]:
       """Perform comprehensive analysis of image contents."""
       results = {}


       # Generate basic caption
       results["caption"] = self.generate_caption(image_path)


       # Detect objects
       objects = self.detect_objects(image_path)
       if objects:
           # Simplify the object list for easier use
           object_list = {}
           for obj in objects:
               label = obj['label']
               if label in object_list:
                   object_list[label] += 1
               else:
                   object_list[label] = 1


           object_summary = ", ".join([f"{count} {label}{'s' if count > 1 else ''}"
                                       for label, count in object_list.items()])
           results["objects_detected"] = object_summary


           # Get bounding box information for major objects
           major_objects = [obj for obj in objects if obj['score'] > 0.5]
           if major_objects:
               # Find central object (closest to image center)
               try:
                   image = Image.open(image_path)
                   image_center = (image.width / 2, image.height / 2)


                   def distance_to_center(obj):
                       box = obj['box']
                       # Ensure keys exist and are numeric
                       xmin = box.get('xmin', 0)
                       ymin = box.get('ymin', 0)
                       xmax = box.get('xmax', 0) # Use xmax, ymax with DETR output
                       ymax = box.get('ymax', 0)


                       obj_center = ((xmin + xmax) / 2, (ymin + ymax) / 2)
                       return ((obj_center[0] - image_center[0]) ** 2 +
                               (obj_center[1] - image_center[1]) ** 2) ** 0.5


                   major_objects.sort(key=distance_to_center)
                   results["central_object"] = major_objects[0]['label']
               except Exception as e:
                   logger.error(f"Error finding central object: {e}")


       # Classify scene
       scene_results = self.classify_scene(image_path)
       if scene_results:
           results["scene_type"] = scene_results[0]['label']


           # Get top 3 scene categories if available
           if len(scene_results) >= 3:
               scene_categories = [f"{r['label']} ({r['score']:.2f})" for r in scene_results[:3]]
               results["scene_categories"] = ", ".join(scene_categories)


       # Extract text from image
       image_text = self.extract_text_from_image(image_path)
       if image_text:
           results["text_in_image"] = image_text


       # Analyze color palette
       colors = self.analyze_color_palette(image_path)
       if colors:
           results["color_palette"] = ", ".join(colors)


       logger.info(f"Completed comprehensive image analysis with {len(results)} attributes")
       return results




class EnhancedRetriever:
   """Enhanced retrieval system using hybrid search and advanced re-ranking."""


   def __init__(self):
       # Initialize embedding model
       self.embedding_model = HuggingFaceEmbeddings(
           model_name="sentence-transformers/all-mpnet-base-v2"
       )


       # Initialize re-ranker with multiple options for fallback
       try:
           # Try to load a more powerful re-ranker first
           self.reranker = CrossEncoder('cross-encoder/ms-marco-MiniLM-L-12-v2')
           logger.info("Initialized enhanced cross-encoder reranker (L-12)")
       except Exception as e:
           logger.warning(f"Failed to load enhanced reranker: {e}")
           try:
               # Fall back to simpler re-ranker
               self.reranker = CrossEncoder('cross-encoder/ms-marco-MiniLM-L-6-v2')
               logger.info("Initialized standard cross-encoder reranker (L-6)")
           except Exception as e2:
               logger.error(f"Failed to load fallback reranker: {e2}")
               self.reranker = None


       # Initialize query expansion with multiple options
       self.query_expander_chain = self._initialize_query_expander()


       # Initialize relevance classifier
       self.relevance_classifier = self._initialize_relevance_classifier()


   def _initialize_query_expander(self) -> Optional[LLMChain]:
       """Initialize the query expansion chain with fallbacks."""
       try:
           # Try T5-base first for better expansion
           query_expander_model = HuggingFaceHub(
               repo_id="google/flan-t5-base",
               model_kwargs={"temperature": 0.7, "max_length": 100}
           )


           query_expansion_prompt = PromptTemplate(
               input_variables=["query"],
               template="""Rewrite the following query to improve document retrieval.
               Add synonyms, related concepts, and expand abbreviations:


               Query: {query}


               Expanded Query:"""
           )


           return LLMChain(
               llm=query_expander_model,
               prompt=query_expansion_prompt
           )


       except Exception as e:
           logger.warning(f"Failed to load primary query expander: {e}")


           try:
               # Fall back to smaller model
               query_expander_model = HuggingFaceHub(
                   repo_id="google/flan-t5-small",
                   model_kwargs={"temperature": 0.5, "max_length": 100}
               )


               query_expansion_prompt = PromptTemplate(
                   input_variables=["query"],
                   template="Expand this query with related terms: {query}"
               )


               return LLMChain(
                   llm=query_expander_model,
                   prompt=query_expansion_prompt
               )


           except Exception as e2:
               logger.error(f"Failed to load fallback query expander: {e2}")
               return None


   def _initialize_relevance_classifier(self) -> Optional[Any]:
       """Initialize a classifier to determine relevance of retrieved documents."""
       try:
           from sentence_transformers import SentenceTransformer, util


           model = SentenceTransformer('all-MiniLM-L6-v2')
           logger.info("Initialized semantic similarity model for relevance classification")
           return model
       except Exception as e:
           logger.warning(f"Failed to load relevance classifier: {e}")
           return None


   def setup_vectorstore(self, chunks: List[Any]) -> Any:
       """Set up the vector store with document chunks."""
       start_time = time.time()
       vectorstore = FAISS.from_documents(documents=chunks, embedding=self.embedding_model)
       end_time = time.time()
       logger.info(f"Vector store created in {end_time - start_time:.2f} seconds")
       return vectorstore


   def expand_query(self, query: str) -> str:
       """Expand the query to improve retrieval with better synonyms and related concepts."""
       if not self.query_expander_chain:
           return query


       try:
           # The run method of LLMChain returns the generated text directly
           expanded_query = self.query_expander_chain.run(query)


           # Clean up potential extra spaces or unwanted characters
           expanded_query = expanded_query.strip()


           # Check if expansion is actually useful
           # Consider if expanded query contains substantially new information
           if len(expanded_query.split()) <= len(query.split()) * 0.8 or expanded_query.lower().strip() == query.lower().strip():
                logger.info(f"Query expansion did not yield substantial improvement, using original query: '{query}'")
                return query


           # Ensure expansion isn't too long (arbitrary limit, can be adjusted)
           if len(expanded_query.split()) > 50:
               expanded_query = " ".join(expanded_query.split()[:50])
               logger.warning(f"Truncated expanded query to 50 words.")




           logger.info(f"Expanded query: '{query}' -> '{expanded_query}'")
           return expanded_query
       except Exception as e:
           logger.error(f"Error expanding query: {e}")
           return query


   def setup_keyword_search(self, chunks: List[Any]) -> Tuple[TfidfVectorizer, Any]:
       """Set up keyword search using TF-IDF with improved parameters."""
       texts = [chunk.page_content for chunk in chunks]


       # Enhanced TF-IDF with better parameters
       # Use NLTK stopwords if available, otherwise default to English
       try:
           stopwords = nltk.corpus.stopwords.words('english')
       except LookupError:
           logger.warning("NLTK stopwords not found. Using default English stopwords.")
           stopwords = ['i', 'me', 'my', 'myself', 'we', 'our', 'ours', 'ourselves', 'you', "you're", "you've", "you'll", "you'd", 'your', 'yours', 'yourself', 'yourselves', 'he', 'him', 'his', 'himself', 'she', "she's", 'her', 'hers', 'herself', 'it', "it's", 'its', 'itself', 'they', 'them', 'their', 'theirs', 'themselves', 'what', 'which', 'who', 'whom', 'this', 'that', "that'll", 'these', 'those', 'am', 'is', 'are', 'was', 'were', 'be', 'been', 'being', 'have', 'has', 'had', 'having', 'do', 'does', 'did', 'doing', 'a', 'an', 'the', 'and', 'but', 'if', 'or', 'because', 'as', 'until', 'while', 'of', 'at', 'by', 'for', 'with', 'about', 'against', 'between', 'into', 'through', 'during', 'before', 'after', 'above', 'below', 'to', 'from', 'up', 'down', 'in', 'out', 'on', 'off', 'over', 'under', 'again', 'further', 'then', 'once', 'here', 'there', 'when', 'where', 'why', 'how', 'all', 'any', 'both', 'each', 'few', 'more', 'most', 'other', 'some', 'such', 'no', 'nor', 'not', 'only', 'own', 'same', 'so', 'than', 'too', 'very', 's', 't', 'can', 'will', 'just', 'don', "don't", 'should', "should've", 'now', 'd', 'll', 'm', 'o', 're', 've', 'y', 'ain', 'aren', "aren't", 'couldn', "couldn't", 'didn', "didn't", 'doesn', "doesn't", 'hadn', "hadn't", 'hasn', "hasn't", 'haven', "haven't", 'isn', "isn't", 'ma', 'mightn', "mightn't", 'mustn', "mustn't", 'needn', "needn't", 'shan', "shan't", 'shouldn', "shouldn't", 'wasn', "wasn't", 'weren', "weren't", 'won', "won't", 'wouldn', "wouldn't"]




       vectorizer = TfidfVectorizer(
           stop_words=stopwords,
           lowercase=True,
           ngram_range=(1, 2), # Consider bi-grams
           max_features=5000 # Limit feature count
           )
       tfidf_matrix = vectorizer.fit_transform(texts)


       logger.info(f"TF-IDF vectorizer trained with {len(vectorizer.get_feature_names_out())} features")


       return vectorizer, tfidf_matrix






def main():
   parser = argparse.ArgumentParser(description="Multimodal RAG Processor")
   parser.add_argument("-i", "--image", required=False, help="Path to input image")
   parser.add_argument("-g", "--guidelines", required=True, help="Path to guidelines PDF")
   parser.add_argument("-m", "--model", required=True, help="Claude model to use")
   parser.add_argument("-q", "--query", required=False, help="Optional custom query")
   args = parser.parse_args()


   load_dotenv()


   processor = DocumentProcessor()
   splitter = AdaptiveTextSplitter()
   retriever = EnhancedRetriever()
   image_processor = ImageProcessor()


   logger.info("Loading and processing documents...")
   docs = processor.load_document(args.guidelines)
   chunks = splitter.split_documents(docs, doc_type="pdf")


   vectorstore = retriever.setup_vectorstore(chunks)


   query = args.query or USER_PROMPT
   query = retriever.expand_query(query)


   similar_chunks = vectorstore.similarity_search(query, k=6)
   context_sections = [f"[SECTION {i + 1}]\n{doc.page_content}" for i, doc in enumerate(similar_chunks)]
   context = "\n\n".join(context_sections)


   # Encode the guideline document
   with open(args.guidelines, "rb") as f:
       b64_guidelines = base64.b64encode(f.read()).decode("utf-8")
   guidelines_media_type = "application/pdf"


   image_data, image_media_type = None, None
   if args.image:
       with open(args.image, "rb") as f:
           image_data = base64.b64encode(f.read()).decode("utf-8")
       image_media_type = "image/png" if args.image.endswith(".png") else "image/jpeg"


   rag_prompt = f"""
   {USER_PROMPT}


   Here are some relevant sections from the document that might help answer the question:


   {context}


   Please analyze the image in relation to the IKEA brand guidelines provided in these sections.
   Identify any violations of the guidelines in the image and provide specific references to the relevant guideline sections.
   Structure your response as a JSON object listing each violation, the guideline reference, and suggested fixes.
   """


   messages = [
       {
           "role": "user",
           "content": [
               {"type": "text", "text": rag_prompt},
               {"type": "document", "source": {"type": "base64", "media_type": guidelines_media_type, "data": b64_guidelines}}
           ]
       }
   ]


   if image_data:
       messages[0]["content"].append({
           "type": "image",
           "source": {
               "type": "base64",
               "media_type": image_media_type,
               "data": image_data
           }
       })


   client = Anthropic(api_key=os.environ.get("ANTHROPIC_API_KEY"))


   logger.info("Sending request to Claude...")
   try:
       response = client.messages.create(
           model=args.model,
           system=SYSTEM_MESSAGE,
           messages=messages,
           max_tokens=4096
       )
       print(response.content[0].text)
   except Exception as e:
       logger.error(f"Claude API call failed: {e}")




if __name__ == "__main__":
   main()

